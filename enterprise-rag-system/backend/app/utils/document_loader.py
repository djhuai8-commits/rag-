"""
多格式文档解析工具
支持：PDF、Word（docx）、Excel（xlsx/xls）、TXT、Markdown
"""
import os
import hashlib
import pandas as pd
from pathlib import Path
from typing import List
from langchain.schema import Document


def load_document(file_path: str) -> List[Document]:
    """根据文件扩展名自动选择解析器"""
    ext = Path(file_path).suffix.lower().lstrip(".")
    loaders = {
        "pdf":  _load_pdf,
        "docx": _load_docx,
        "doc":  _load_docx,
        "xlsx": _load_excel,
        "xls":  _load_excel,
        "txt":  _load_text,
        "md":   _load_text,
    }
    loader_fn = loaders.get(ext)
    if loader_fn is None:
        raise ValueError(f"不支持的文件格式：.{ext}")
    docs = loader_fn(file_path)
    # 注入通用元数据
    file_hash = _calc_md5(file_path)
    for doc in docs:
        doc.metadata.setdefault("source", os.path.basename(file_path))
        doc.metadata.setdefault("doc_type", ext)
        doc.metadata["file_hash"] = file_hash
    return docs


def _load_pdf(file_path: str) -> List[Document]:
    """PDF 解析：优先 PyMuPDF，扫描件自动切换 Unstructured OCR"""
    try:
        from langchain_community.document_loaders import PyMuPDFLoader
        loader = PyMuPDFLoader(file_path)
        docs = loader.load()
        # 文本量过少时认为是扫描件，切换 OCR
        total_text = "".join(d.page_content for d in docs)
        if len(total_text.strip()) < 100:
            raise ValueError("文本量不足，疑似扫描件")
        return docs
    except Exception:
        from langchain_community.document_loaders import UnstructuredPDFLoader
        loader = UnstructuredPDFLoader(
            file_path,
            strategy="hi_res",
            extract_images_in_pdf=False,
            infer_table_structure=True,
        )
        return loader.load()


def _load_docx(file_path: str) -> List[Document]:
    """Word 文档解析，保留段落层级和表格"""
    try:
        import docx
        doc = docx.Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        # 处理表格
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append(" | ".join(cell.text.strip() for cell in row.cells))
            if rows:
                paragraphs.append("\n".join(rows))
        content = "\n\n".join(paragraphs)
        return [Document(page_content=content, metadata={"source": file_path})]
    except Exception:
        from langchain_community.document_loaders import Docx2txtLoader
        loader = Docx2txtLoader(file_path)
        return loader.load()


def _load_excel(file_path: str) -> List[Document]:
    """Excel 解析，每个 Sheet 生成一个 Document"""
    docs = []
    xls = pd.ExcelFile(file_path)
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name, dtype=str).fillna("")
        content = df.to_markdown(index=False)
        if content and content.strip():
            docs.append(Document(
                page_content=f"# {sheet_name}\n\n{content}",
                metadata={"source": file_path, "sheet": sheet_name},
            ))
    return docs if docs else [Document(page_content="", metadata={"source": file_path})]


def _load_text(file_path: str) -> List[Document]:
    """TXT / Markdown 纯文本解析"""
    for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read()
            return [Document(page_content=content, metadata={"source": file_path})]
        except UnicodeDecodeError:
            continue
    raise ValueError(f"无法识别文件编码：{file_path}")


def _calc_md5(file_path: str) -> str:
    h = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()
